"""
Phase B: 정답데이터2 521건 전체 카드화 (사용자 통찰 + 김보민 쪽지 직접 적용)

설계 원칙 (Phase A 인사이트 반영):
1. EVAL_ONLY 라벨 폐지 — TOBE 채워진 521건 모두 카드화
2. NEGATIVE만 분리 (협의=기존유지·삭제 또는 피드백=미수용·원안유지)
3. 시트별 컬럼 매핑 정확 적용 (1·2·3·4차)
4. 시트 4 customer_confirm·additional_suggest 메타 추가
5. 시트 2 wirelink_memo 메타 추가
6. 시트 3 msg_code 없음 → msg_name·시트_R번호 사용
7. process 도메인 신규 (시트 3 전용)
8. v3 마커 의미 정정 — 외주업체 추가 개선 (학습 자료)

출력: output/정답데이터2/ 폴더에 도메인별 카드 docx (8 도메인 + 1 NEGATIVE = 9 파일)
        그 후 split_cards_v2.py 실행으로 3000자 분할
"""

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET
from openpyxl import load_workbook
from collections import defaultdict
from docx import Document

XLSX_PATH = Path(r"C:\dev\guide\정답데이터2\정답데이터2.xlsx")
OUTPUT_DIR = Path(r"C:\dev\guide\output\정답데이터2")

# 시트별 컬럼 매핑
SHEET_CFG = {
    "1차_국내해외주식채권_피드백반영": {
        "msg_code": 2, "msg_name": 3, "asis": 8, "tobe": 9,
        "feedback": 10, "team": 16, "review": 18, "agree": 21,
    },
    "2차_파생 등_피드백반영": {
        "msg_code": 2, "msg_name": 3, "asis": 8, "tobe": 9,
        "team": 14, "review": 16, "agree": 18, "feedback": 19,
        "wirelink_memo": 20,
    },
    "3차_프로세스 외_피드백 반영": {
        "msg_name": 2, "asis": 3, "tobe": 4, "feedback": 5,
        "team": 6, "agree": 8, "memo": 9,
    },
    "4차_ 연금 전체": {
        "msg_code": 2, "msg_name": 3, "asis": 8, "tobe": 9,
        "customer_confirm": 10, "additional_suggest": 11,
        "team": 12, "review": 14, "agree": 16, "feedback": 17,
    },
}

NEGATIVE_AGREE = {"기존유지", "삭제"}
NEGATIVE_FEEDBACK_KW = {"미수용", "원안유지"}
POSITIVE_FEEDBACK_KW = {"수용", "이견없음", "이상없음", "이의없음", "진행요청"}

CTRL_CHAR_RE = re.compile(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]")


def clean(text):
    return CTRL_CHAR_RE.sub("", str(text or ""))


# ============== rich text 색 정보 ==============
NS = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}


def normalize_color(hex_or_indexed):
    if not hex_or_indexed:
        return "DEFAULT"
    s = str(hex_or_indexed).upper()
    if len(s) == 8:
        s = s[2:]
    if s in ("FF0000", "C00000", "FF2600", "990000", "B30000"):
        return "RED"
    if s in ("0070C0", "0066CC", "0033CC", "0000FF", "1F4E79", "2E75B6"):
        return "BLUE"
    if s in ("000000", "DEFAULT", "1F1F1F"):
        return "BLACK"
    return s


def parse_rich_text_segments(rt_element):
    segments = []
    for r in rt_element.findall("main:r", NS):
        text = ""
        t = r.find("main:t", NS)
        if t is not None and t.text:
            text = t.text
        color = "DEFAULT"
        rpr = r.find("main:rPr", NS)
        if rpr is not None:
            color_elem = rpr.find("main:color", NS)
            if color_elem is not None:
                rgb = color_elem.get("rgb")
                if rgb:
                    color = normalize_color(rgb)
        if text:
            segments.append({"text": text, "color": color})
    if not segments:
        t = rt_element.find("main:t", NS)
        if t is not None and t.text:
            segments.append({"text": t.text, "color": "DEFAULT"})
    return segments


def load_shared_strings_segments(xlsx_path):
    with zipfile.ZipFile(xlsx_path, "r") as z:
        try:
            with z.open("xl/sharedStrings.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
        except KeyError:
            return {}
    result = {}
    for idx, si in enumerate(root.findall("main:si", NS)):
        result[idx] = parse_rich_text_segments(si)
    return result


def get_sheet_paths(xlsx_path):
    with zipfile.ZipFile(xlsx_path, "r") as z:
        with z.open("xl/workbook.xml") as f:
            wb_root = ET.parse(f).getroot()
        with z.open("xl/_rels/workbook.xml.rels") as f:
            rels_root = ET.parse(f).getroot()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    rid_to_target = {r.get("Id"): r.get("Target") for r in rels_root.findall(f"{{{rels_ns}}}Relationship")}

    name_to_path = {}
    for s in wb_root.findall("main:sheets/main:sheet", NS):
        name = s.get("name")
        r_id = s.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        target = rid_to_target.get(r_id)
        if target:
            path = f"xl/{target}" if not target.startswith("xl/") else target
            name_to_path[name] = path
    return name_to_path


# 시트별 sheetData 캐시
_sheet_cell_cache = {}


def build_sheet_cell_map(xlsx_path, sheet_name, sheet_path):
    """sheet의 모든 셀 (ref → string_index) 매핑 반환"""
    cache_key = (str(xlsx_path), sheet_name)
    if cache_key in _sheet_cell_cache:
        return _sheet_cell_cache[cache_key]
    with zipfile.ZipFile(xlsx_path, "r") as z:
        with z.open(sheet_path) as f:
            sheet_root = ET.parse(f).getroot()
    cell_map = {}
    for c in sheet_root.findall("main:sheetData/main:row/main:c", NS):
        if c.get("t") == "s":
            v = c.find("main:v", NS)
            if v is not None and v.text:
                cell_map[c.get("r")] = int(v.text)
    _sheet_cell_cache[cache_key] = cell_map
    return cell_map


def col_letter(col):
    if col <= 26:
        return chr(ord("A") + col - 1)
    # 27=AA
    first = (col - 1) // 26
    second = (col - 1) % 26
    return chr(ord("A") + first - 1) + chr(ord("A") + second)


def extract_cell_segments(xlsx_path, shared_strings, sheet_paths, sheet_name, row, col, fallback_value):
    sheet_path = sheet_paths.get(sheet_name)
    if not sheet_path:
        if fallback_value:
            return [{"text": str(fallback_value), "color": "DEFAULT"}]
        return []
    cell_map = build_sheet_cell_map(xlsx_path, sheet_name, sheet_path)
    cell_ref = f"{col_letter(col)}{row}"
    idx = cell_map.get(cell_ref)
    if idx is None:
        if fallback_value:
            return [{"text": str(fallback_value), "color": "DEFAULT"}]
        return []
    return shared_strings.get(idx, [])


def render_inline_marked(segments):
    if not segments:
        return ""
    parts = []
    for s in segments:
        text = s["text"]
        color = s["color"]
        if color == "BLUE":
            parts.append(f"<v2>{text}</v2>")
        elif color == "RED":
            parts.append(f"<v3>{text}</v3>")
        else:
            parts.append(text)
    return "".join(parts)


# ============== 라벨·도메인 분류 ==============
def is_negative(agree, feedback):
    a = str(agree or "").strip()
    f = str(feedback or "").strip()
    if a in NEGATIVE_AGREE:
        return True
    if any(kw in a for kw in ["기존유지", "삭제"]):
        return True
    if any(kw in f for kw in NEGATIVE_FEEDBACK_KW):
        return True
    return False


def classify_label(agree, feedback, review):
    if is_negative(agree, feedback):
        return "NEGATIVE"
    a = str(agree or "").strip()
    r = str(review or "").strip()
    f = str(feedback or "").strip()
    if a == "반영" and any(kw in f for kw in POSITIVE_FEEDBACK_KW):
        return "HIGH"
    if a == "반영":
        return "HIGH"
    if r in ("검토완료", "검수완료"):
        return "MEDIUM"
    if any(kw in f for kw in POSITIVE_FEEDBACK_KW):
        return "MEDIUM"
    return "LOW"


TEAM_DOMAIN = {
    "증권운영팀": "credit_loan", "신용관리팀": "credit_loan",
    "파생상품팀": "derivatives", "파생서비스팀": "derivatives",
    "국제거래팀": "derivatives", "CFD팀": "derivatives",
    "랩어카운트팀": "product", "상품서비스팀": "product",
    "신상품팀": "product", "투자상품팀": "product",
    "결제팀": "settlement", "결제서비스팀": "settlement",
    "연금업무개발팀": "pension", "퇴직연금팀": "pension",
    "연금사업본부": "pension", "연금영업팀": "pension",
    "디지털마케팅팀": "marketing", "브랜드마케팅팀": "marketing",
    "채널서비스팀": "process", "디지털기획팀": "process",
    "개인정보보호팀": "process", "고객만족팀": "process",
}


def classify_domain(team, msg_name, sheet_name):
    if "3차_프로세스" in sheet_name:
        return "process"
    if "4차_ 연금" in sheet_name:
        return "pension"
    team_str = str(team or "").strip()
    msg_str = str(msg_name or "").strip()
    name_lower = msg_str.lower()
    if any(kw in msg_str for kw in ["연금", "IRP", "퇴직"]) or "irp" in name_lower:
        return "pension"
    if any(kw in msg_str for kw in ["이벤트", "마케팅", "광고", "혜택"]):
        return "marketing"
    for key, dom in TEAM_DOMAIN.items():
        if key in team_str:
            return dom
    return "misc"


# ============== 카드 추출 ==============
def extract_cards(xlsx_path):
    print("[1/3] xlsx sharedStrings 파싱...")
    shared_strings = load_shared_strings_segments(xlsx_path)
    print(f"  shared strings: {len(shared_strings)}개")

    print("[2/3] 시트 경로 매핑...")
    sheet_paths = get_sheet_paths(xlsx_path)
    print(f"  시트 {len(sheet_paths)}개")

    print("[3/3] 카드 추출...")
    wb = load_workbook(xlsx_path, data_only=True)
    all_cards = []

    for sheet_name, cfg in SHEET_CFG.items():
        if sheet_name not in wb.sheetnames:
            print(f"  WARN: {sheet_name} 없음")
            continue
        ws = wb[sheet_name]
        sheet_cards = 0
        for r in range(3, ws.max_row + 1):
            tobe_val = ws.cell(row=r, column=cfg["tobe"]).value
            if not tobe_val or not str(tobe_val).strip():
                continue

            asis_val = ws.cell(row=r, column=cfg["asis"]).value if cfg.get("asis") else None
            asis_segs = extract_cell_segments(xlsx_path, shared_strings, sheet_paths, sheet_name, r, cfg["asis"], asis_val) if cfg.get("asis") else []
            tobe_segs = extract_cell_segments(xlsx_path, shared_strings, sheet_paths, sheet_name, r, cfg["tobe"], tobe_val)

            def get(key):
                col = cfg.get(key)
                if not col:
                    return None
                return ws.cell(row=r, column=col).value

            agree = get("agree")
            feedback = get("feedback")
            review = get("review")
            label = classify_label(agree, feedback, review)
            domain = classify_domain(get("team"), get("msg_name"), sheet_name)

            msg_code = str(get("msg_code") or "").strip()
            msg_name = str(get("msg_name") or "").strip()
            if not msg_code:
                msg_code = f"{sheet_name[:5]}_R{r}"  # fallback ID

            card = {
                "sheet": sheet_name, "row": r,
                "msg_code": msg_code,
                "msg_name": msg_name,
                "team": str(get("team") or "").strip(),
                "asis_marked": render_inline_marked(asis_segs),
                "tobe_marked": render_inline_marked(tobe_segs),
                "agree": str(agree or "").strip(),
                "review": str(review or "").strip(),
                "feedback": str(feedback or "").strip(),
                "label": label,
                "domain": domain,
            }
            if "4차" in sheet_name:
                card["customer_confirm"] = str(get("customer_confirm") or "").strip()
                card["additional_suggest"] = str(get("additional_suggest") or "").strip()
            if "2차" in sheet_name:
                card["wirelink_memo"] = str(get("wirelink_memo") or "").strip()

            all_cards.append(card)
            sheet_cards += 1
        print(f"  {sheet_name}: {sheet_cards}건")

    print(f"\n총 추출: {len(all_cards)}건")
    return all_cards


# ============== msg_code 중복 제거 ==============
def deduplicate_by_msg_code(cards):
    label_priority = {"HIGH": 1, "MEDIUM": 2, "LOW": 3, "NEGATIVE": 4}
    by_code = defaultdict(list)
    no_code = []
    for c in cards:
        # fallback ID (시트_R번호)는 중복 제거 대상 아님
        if c["msg_code"] and not c["msg_code"].startswith(("1차_R", "2차_R", "3차_R", "4차_R")):
            by_code[c["msg_code"]].append(c)
        else:
            no_code.append(c)
    deduped = []
    for code, cs in by_code.items():
        cs_sorted = sorted(cs, key=lambda c: label_priority.get(c["label"], 99))
        deduped.append(cs_sorted[0])
    return deduped + no_code


# ============== 카드 docx 생성 ==============
def make_card_text_lines(card):
    lines = []
    lines.append("━" * 67)
    lines.append(f"### 사례 [{card['msg_code'][:30]}] {card.get('msg_name', '')[:60]}")
    lines.append("━" * 67)
    lines.append("")
    lines.append(f"🎯 이 사례는 {card.get('team') or '담당팀 미정'}의 '{card.get('msg_name', '')}' 메시지를 검수 통과 형태로 개선한 정답지입니다.")
    lines.append("")
    lines.append("━━━ ASIS (직원 작성 원본) ━━━")
    for line in card.get("asis_marked", "").split("\n"):
        lines.append(line)
    lines.append("")
    lines.append("━━━ TOBE (검수자 개선안 — 인라인 마커) ━━━")
    for line in card.get("tobe_marked", "").split("\n"):
        lines.append(line)
    lines.append("")
    lines.append("━━━ 메타데이터 ━━━")
    lines.append(f"- 신뢰도 라벨: {card.get('label', 'LOW')}")
    lines.append(f"- 도메인: {card.get('domain', 'misc')}")
    if card.get("team"):
        lines.append(f"- 담당팀: {card['team']}")
    lines.append(f"- 메시지 코드: {card['msg_code']}")
    if card.get("msg_name"):
        lines.append(f"- 메시지 이름: {card['msg_name']}")
    if card.get("agree"):
        lines.append(f"- 협의결과: {card['agree']}")
    if card.get("review"):
        lines.append(f"- 검토결과: {card['review']}")
    if card.get("feedback"):
        feedback_first = str(card["feedback"]).split("\n")[0][:80]
        lines.append(f"- 담당팀 피드백 첫 라인: {feedback_first}")
    if card.get("customer_confirm"):
        cc = str(card["customer_confirm"]).replace("\n", " ")[:120]
        lines.append(f"- 고객사 컨펌의견: {cc}")
    if card.get("additional_suggest"):
        as_ = str(card["additional_suggest"]).replace("\n", " ")[:120]
        lines.append(f"- 와이어링크 추가 제안(02/25): {as_}")
    if card.get("wirelink_memo"):
        wm = str(card["wirelink_memo"]).replace("\n", " ")[:120]
        lines.append(f"- 와이어링크 메모: {wm}")
    lines.append(f"- 원천: 정답데이터2 / {card.get('sheet')} R{card.get('row')}")
    return lines


def write_docx(cards, output_path, header_title, header_subtitle):
    doc = Document()
    doc.add_heading(clean(header_title), level=1)
    doc.add_paragraph(clean(header_subtitle))
    doc.add_paragraph(f"총 {len(cards)}건의 사례가 포함되어 있습니다.")
    doc.add_paragraph("")
    doc.add_paragraph("━" * 67)
    doc.add_paragraph("【마커 해석 가이드 (전체 카드 공통)】")
    doc.add_paragraph("- 마커 없는 텍스트 = TOBE 검수자 1차 초안 (신뢰도 중)")
    doc.add_paragraph("- <v2>...</v2> = 현업 합의 (와이어링크 1교, 신뢰도 고)")
    doc.add_paragraph("- <v3>...</v3> = 외주업체(와이어링크) 추가 개선 또는 다영 2교 (가이드 부합도 더 높음, 학습 자료)")
    doc.add_paragraph("━" * 67)
    doc.add_paragraph("")

    for card in cards:
        for line in make_card_text_lines(card):
            doc.add_paragraph(clean(line))

    doc.save(str(output_path))


def main():
    print("=== Phase B: 정답데이터2 521건 전체 카드화 ===\n")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # 기존 분할본·원본 docx 모두 삭제 (clean rebuild)
    print("[기존 docx 삭제 중...]")
    deleted = 0
    for f in OUTPUT_DIR.glob("cx_*.docx"):
        f.unlink()
        deleted += 1
    print(f"  {deleted}개 삭제\n")

    # 카드 추출
    cards = extract_cards(XLSX_PATH)

    # 중복 제거
    cards_dedup = deduplicate_by_msg_code(cards)
    print(f"\nmsg_code 중복 제거 후: {len(cards_dedup)}건 (제거 {len(cards) - len(cards_dedup)}건)")

    # 라벨 분포
    by_label = defaultdict(list)
    for c in cards_dedup:
        by_label[c["label"]].append(c)
    print(f"라벨 분포: HIGH={len(by_label['HIGH'])}, MEDIUM={len(by_label['MEDIUM'])}, "
          f"LOW={len(by_label['LOW'])}, NEGATIVE={len(by_label['NEGATIVE'])}")

    # 카드용 (NEGATIVE 제외) 도메인 그룹핑
    cards_for_db = by_label["HIGH"] + by_label["MEDIUM"] + by_label["LOW"]
    grouped = defaultdict(list)
    for c in cards_for_db:
        grouped[c["domain"]].append(c)

    print(f"\n도메인별 분포:")
    for domain, cs in sorted(grouped.items(), key=lambda x: -len(x[1])):
        print(f"  {domain}: {len(cs)}건")

    # 도메인별 docx 생성
    domain_to_filename = {
        "credit_loan": "cx_goldens2_credit_loan.docx",
        "derivatives": "cx_goldens2_derivatives.docx",
        "product": "cx_goldens2_product.docx",
        "settlement": "cx_goldens2_settlement.docx",
        "pension": "cx_goldens2_pension.docx",
        "marketing": "cx_goldens2_marketing.docx",
        "process": "cx_goldens2_process.docx",  # 신규 도메인
        "misc": "cx_goldens2_misc.docx",
    }

    print(f"\n[카드 docx 생성]")
    for domain, filename in domain_to_filename.items():
        cs = grouped.get(domain, [])
        if not cs:
            continue
        path = OUTPUT_DIR / filename
        write_docx(cs, path,
                   header_title=f"미래에셋증권 CX 정답지 (2차 개선 - {domain})",
                   header_subtitle=f"회사 공식 2차 메시지 개선 521건 중 {domain} 도메인의 검수 통과 사례 모음. "
                                   f"각 사례는 ASIS(원문)와 TOBE(검수자 개선안)를 비교하며, "
                                   f"TOBE의 색별 신뢰도(검정/파랑/빨강)는 인라인 마커(<v2>/<v3>)로 표현됩니다. "
                                   f"v2=현업 합의(와이어링크 1교), v3=외주업체 추가 개선 또는 다영 2교(가이드 부합도 더 높음).")
        print(f"  ✅ {filename}: {len(cs)}건")

    # NEGATIVE
    if by_label["NEGATIVE"]:
        path = OUTPUT_DIR / "cx_exceptions_misuyong.docx"
        write_docx(by_label["NEGATIVE"], path,
                   header_title="미래에셋증권 CX 예외 사례 (가이드 룰 적용 자제)",
                   header_subtitle="협의결과 = '기존유지/삭제' 또는 담당팀 피드백 = '미수용/원안유지' 사례. "
                                   "이 메시지 코드/패턴이 등장하면 가이드 룰을 강하게 적용하지 말고 원안에 가깝게 유지할 것.")
        print(f"  ✅ cx_exceptions_misuyong.docx: {len(by_label['NEGATIVE'])}건")

    print("\n=== Phase B 완료 ===")
    print(f"카드 총합: HIGH·MEDIUM·LOW {len(cards_for_db)}건 + NEGATIVE {len(by_label['NEGATIVE'])}건 = {len(cards_dedup)}건")
    print(f"다음 단계: split_cards_v2.py 실행으로 3000자 분할")


if __name__ == "__main__":
    main()
