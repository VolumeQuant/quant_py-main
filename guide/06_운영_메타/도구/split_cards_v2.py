"""
정답데이터2 카드 docx 3000자 단위 분할

문제: 카드 docx 9600~40700자 → vector DB 임베딩 시 오류·삽입 실패 위험.
해결: 사례 단위로 묶어 3000자 이내 분할 (헤더·마커 가이드는 모든 분할 파일에 복사).

정답데이터1은 3000~6500자라 분할 불필요.
"""

import re
from docx import Document
from pathlib import Path

INPUT_DIR = Path(r"C:\dev\guide\output\정답데이터2")
OUTPUT_DIR = Path(r"C:\dev\guide\output\정답데이터2")  # 같은 폴더 — 원본 삭제 후 분할본 저장
TARGET_CHARS = 2200  # 헤더 ~600자 + 안전 마진 200자 차감 → 단일 사례 큰 경우 외 모두 3000자 이내

CTRL_CHAR_RE = re.compile(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]")


def clean(text):
    return CTRL_CHAR_RE.sub("", text)


def split_into_cases(lines):
    """헤더 + 사례 리스트로 분리 (사례 시작 = '### 사례' 줄)"""
    header = []
    cases = []
    current = None
    for line in lines:
        if line.strip().startswith("### 사례"):
            if current is not None:
                cases.append(current)
            current = [line]
        elif current is None:
            header.append(line)
        else:
            current.append(line)
    if current is not None:
        cases.append(current)
    return header, cases


def make_chunk_docx(header_lines, cases, output_path):
    doc = Document()
    for line in header_lines:
        doc.add_paragraph(clean(line))
    for case in cases:
        for line in case:
            doc.add_paragraph(clean(line))
    doc.save(str(output_path))


def split_docx(input_path, output_dir, target_chars=TARGET_CHARS):
    doc = Document(str(input_path))
    lines = [p.text for p in doc.paragraphs]

    header, cases = split_into_cases(lines)

    # 3000자 단위로 묶음
    chunks = []
    cur_chunk = []
    cur_chars = 0
    for case in cases:
        case_chars = sum(len(line) for line in case)
        if cur_chunk and cur_chars + case_chars > target_chars:
            chunks.append(cur_chunk)
            cur_chunk = []
            cur_chars = 0
        cur_chunk.append(case)
        cur_chars += case_chars
    if cur_chunk:
        chunks.append(cur_chunk)

    # 분할 docx 저장
    base_name = input_path.stem
    out_files = []
    for i, chunk in enumerate(chunks, start=1):
        out_path = output_dir / f"{base_name}_p{i:02d}.docx"
        chunk_header = list(header) + [f"(분할 {i}/{len(chunks)} - {len(chunk)}건)", ""]
        make_chunk_docx(chunk_header, chunk, out_path)
        chunk_chars = sum(sum(len(l) for l in c) for c in chunk)
        out_files.append((out_path.name, len(chunk), chunk_chars))
    return out_files


def main():
    print("=== 정답데이터2 카드 3000자 단위 분할 ===\n")

    targets = sorted(INPUT_DIR.glob("cx_*.docx"))
    print(f"분할 대상: {len(targets)}개 docx\n")

    total_new_files = 0
    for target in targets:
        # 원본 사이즈 확인
        doc = Document(str(target))
        original_chars = sum(len(p.text) for p in doc.paragraphs)
        if original_chars <= TARGET_CHARS:
            print(f"[SKIP] {target.name}: {original_chars}자 (3000자 이내, 분할 불필요)")
            continue

        out_files = split_docx(target, OUTPUT_DIR)
        print(f"[분할] {target.name}: {original_chars}자 → {len(out_files)}개 파일")
        for name, n_cases, n_chars in out_files:
            print(f"        {name}: {n_cases}건 / {n_chars}자")
        total_new_files += len(out_files)

        # 원본 삭제 (분할본만 남김)
        target.unlink()

    print(f"\n=== 분할 완료: {total_new_files}개 파일 생성 ===")


if __name__ == "__main__":
    main()
