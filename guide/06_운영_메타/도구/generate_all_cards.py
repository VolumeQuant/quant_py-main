"""
4단계 통합 변환 도구: 정답데이터2 + 정답데이터1 → 카드 docx 일괄 생성

출력:
  C:/dev/guide/정답데이터2/
    - cx_goldens2_derivatives.docx
    - cx_goldens2_product.docx
    - cx_goldens2_credit_loan.docx
    - cx_goldens2_settlement.docx
    - cx_goldens2_misc.docx
    - cx_goldens2_marketing.docx
    - cx_goldens2_pension.docx
    - cx_exceptions_misuyong.docx (NEGATIVE)

  C:/dev/guide/정답데이터1/
    - cx_goldens_01a~f.docx (덮어쓰기, v5.21 양식, 신뢰도 HIGH 일괄)

처리 흐름:
  1. 정답데이터2 1·2·3·4차 시트 모두 추출
  2. 5축 분류 + 신뢰도 라벨링
  3. msg_code 중복 제거 (HIGH > MEDIUM > LOW > NEGATIVE > EVAL_ONLY)
  4. 도메인별 그룹핑
  5. 카드 텍스트 생성 (인라인 마커 <v2>/<v3>)
  6. python-docx로 docx 생성
  7. 정답데이터1 31건도 같은 양식으로 변환 (HIGH 일괄, 마커 없음)
"""

import zipfile
import re
import xml.etree.ElementTree as ET
import sys
from collections import defaultdict
from pathlib import Path

import openpyxl
from docx import Document

sys.path.insert(0, r"C:\dev\guide\tools")
from classify_5axis_sample import (
    parse_shared_strings, get_sheet_xml_paths, parse_sheet_cells,
    get_rich_text, normalize_color, color_distribution, first_line,
    classify_team_to_domain, classify_pension_sub, assign_5axis_label,
    SHEET_CFG, XLSX_PATH,
)

# ===== 3차 시트 컬럼 매핑 추가 =====
# 3차는 컬럼 구조가 다름 (10컬럼)
# C1: no, C2: 문자메시지상세명, C3: ASIS, C4: TOBE,
# C5: 담당팀 피드백, C6: 2025년 담당 조직명, C7: 비고,
# C8: 협의결과, C9: 메모, C10: ?
SHEET_CFG_3RD = {
    "asis": 3, "tobe": 4, "feedback": 5, "team": 6,
    "review": None,  # 3차는 검토결과 컬럼 없음
    "agree": 8, "code": None, "name": 2, "remark": 7,
}

XLSX_GOLDENS1 = r"C:\dev\guide\정답데이터1\정답데이터1.xlsx"


def merge_consecutive_segments(segments):
    """연속된 같은 색 segment 합침."""
    merged = []
    for s in segments:
        color = normalize_color(s["color"])
        if merged and merged[-1]["color"] == color:
            merged[-1]["text"] += s["text"]
        else:
            merged.append({"color": color, "text": s["text"]})
    return merged


def render_inline_marked(segments):
    """통합본 텍스트에 인라인 색 마커 적용."""
    if not segments:
        return ""
    merged = merge_consecutive_segments(segments)
    parts = []
    for s in merged:
        c = s["color"]
        text = s["text"]
        if c == "BLUE":
            parts.append(f"<v2>{text}</v2>")
        elif c == "RED":
            parts.append(f"<v3>{text}</v3>")
        else:
            parts.append(text)
    return "".join(parts)


def make_card_text(card_data):
    """카드 1개의 plain text 생성."""
    msg_code = card_data["msg_code"] or "(코드없음)"
    msg_name = card_data["msg_name"] or "(이름없음)"
    team = card_data["team"] or ""
    domain = card_data["domain"]
    label = card_data["label"]
    asis_marked = card_data["asis_marked"]
    tobe_marked = card_data["tobe_marked"]
    agree = card_data["agree"]
    review = card_data["review"]
    feedback_first = card_data["feedback_first"]
    sheet_name = card_data["sheet"]
    row = card_data["row"]
    source = card_data.get("source", "정답데이터2")

    # HyDE 한 줄
    if team:
        hyde = f"이 사례는 {team}의 '{msg_name}' 메시지를 검수 통과 형태로 개선한 정답지입니다."
    else:
        hyde = f"이 사례는 '{msg_name}' 메시지를 검수 통과 형태로 개선한 정답지입니다."

    lines = []
    lines.append("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    lines.append(f"### 사례 [{msg_code}] {msg_name}")
    lines.append("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    lines.append("")
    lines.append(f"🎯 {hyde}")
    lines.append("")
    lines.append("━━━ ASIS (직원이 작성한 원본) ━━━")
    lines.append(asis_marked)
    lines.append("")
    lines.append("━━━ TOBE (검수자 개선안 — 인라인 마커) ━━━")
    lines.append(tobe_marked)
    lines.append("")
    lines.append("━━━ 마커 해석 가이드 ━━━")
    lines.append("- 마커 없는 텍스트 = TOBE 검수자 1차 초안 (신뢰도 중)")
    lines.append("- <v2>...</v2>     = 현업 합의 부분 (신뢰도 고 ★ 가장 우선해서 모방)")
    lines.append("- <v3>...</v3>     = 재개선 진행 중 (신뢰도 잠정 — 직접 인용 금지, 방향성만 참조)")
    lines.append("- ASIS의 <v3>...</v3> = 검수자가 '이 부분 빼라'고 표시한 영역")
    lines.append("")
    lines.append("━━━ 메타데이터 ━━━")
    lines.append(f"- 신뢰도 라벨: {label}")
    lines.append(f"- 도메인: {domain}")
    lines.append(f"- 담당팀: {team}")
    lines.append(f"- 메시지 코드: {msg_code}")
    lines.append(f"- 메시지 이름: {msg_name}")
    lines.append(f"- 협의 결과: {agree if agree else '(blank)'}")
    lines.append(f"- 검토 결과: {review if review else '(blank)'}")
    lines.append(f"- 담당팀 피드백 첫 라인: {feedback_first if feedback_first else '(blank)'}")
    lines.append(f"- 원천: {source} / {sheet_name} R{row}")
    lines.append("")

    return "\n".join(lines)


def extract_cards_from_xlsx(xlsx_path, sheets_config_override=None, source_label="정답데이터2"):
    """xlsx에서 모든 카드 데이터 추출."""
    cards = []

    with zipfile.ZipFile(xlsx_path) as z:
        shared_strings = parse_shared_strings(z)
        sheet_map = get_sheet_xml_paths(z)
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)

        sheets_config = sheets_config_override or {
            "1차_국내해외주식채권_피드백반영": SHEET_CFG["1차_국내해외주식채권_피드백반영"],
            "2차_파생 등_피드백반영": SHEET_CFG["2차_파생 등_피드백반영"],
            "3차_프로세스 외_피드백 반영": SHEET_CFG_3RD,
            "4차_ 연금 전체": SHEET_CFG["4차_ 연금 전체"],
        }

        for sn, cfg in sheets_config.items():
            if sn not in sheet_map:
                continue
            cells = parse_sheet_cells(z, sheet_map[sn])
            ws = wb[sn]

            for r in range(3, ws.max_row + 1):
                asis = ws.cell(r, cfg["asis"]).value
                tobe = ws.cell(r, cfg["tobe"]).value
                if not (asis and tobe and len(str(asis)) > 30 and len(str(tobe)) > 30):
                    continue

                # rich text segment 추출 (단, 3차는 inline 단순 텍스트)
                asis_segs = get_rich_text(cells, shared_strings, r, cfg["asis"])
                tobe_segs = get_rich_text(cells, shared_strings, r, cfg["tobe"])

                # 5축 신호
                agree = str(ws.cell(r, cfg["agree"]).value or "").strip() if cfg.get("agree") else ""
                review = str(ws.cell(r, cfg["review"]).value or "").strip() if cfg.get("review") else ""
                feedback = first_line(ws.cell(r, cfg["feedback"]).value if cfg.get("feedback") else None)
                team_raw = ws.cell(r, cfg["team"]).value if cfg.get("team") else None
                team = str(team_raw or "").strip().split("\n")[0] if team_raw else ""
                msg_code = str(ws.cell(r, cfg["code"]).value or "").strip() if cfg.get("code") else ""
                msg_name = str(ws.cell(r, cfg["name"]).value or "").strip() if cfg.get("name") else ""

                # 라벨링
                tobe_color_dist = color_distribution(tobe_segs)
                label = assign_5axis_label(sn, agree, feedback, review, tobe_color_dist)

                # 도메인
                domain = classify_team_to_domain(team)
                if domain == "pension":
                    sub = classify_pension_sub(msg_name)
                    domain_full = f"pension_{sub}" if sub != "misc" else "pension"
                else:
                    domain_full = domain

                cards.append({
                    "sheet": sn, "row": r,
                    "msg_code": msg_code, "msg_name": msg_name, "team": team,
                    "domain": domain_full, "label": label,
                    "asis_marked": render_inline_marked(asis_segs),
                    "tobe_marked": render_inline_marked(tobe_segs),
                    "agree": agree, "review": review, "feedback_first": feedback,
                    "source": source_label,
                })
    return cards


def deduplicate_by_msg_code(cards):
    """msg_code 중복 제거. 신뢰도 HIGH > MEDIUM > LOW > NEGATIVE > EVAL_ONLY 순."""
    label_priority = {"HIGH": 1, "MEDIUM": 2, "LOW": 3, "NEGATIVE": 4, "EVAL_ONLY": 5}

    by_code = defaultdict(list)
    no_code = []
    for c in cards:
        if c["msg_code"]:
            by_code[c["msg_code"]].append(c)
        else:
            no_code.append(c)  # 코드 없는 건 그대로 (3차 시트 일부)

    deduped = []
    for code, cs in by_code.items():
        # 우선순위 정렬 (낮은 priority가 먼저)
        cs_sorted = sorted(cs, key=lambda c: label_priority.get(c["label"], 99))
        deduped.append(cs_sorted[0])

    return deduped + no_code


def domain_grouping(cards, target_domains=None):
    """도메인별로 카드 그룹핑."""
    if target_domains is None:
        target_domains = ["derivatives", "product", "credit_loan", "settlement",
                         "misc", "marketing", "pension", "pension_irp", "pension_dc",
                         "pension_misc"]
    grouped = defaultdict(list)
    for c in cards:
        d = c["domain"]
        # pension 변형은 모두 pension으로 통합 (Phase 7 결정)
        if d.startswith("pension"):
            d = "pension"
        grouped[d].append(c)
    return grouped


def write_docx(cards, output_path, header_title, header_subtitle):
    """카드 리스트를 docx 파일로 저장."""
    doc = Document()
    # 헤더
    doc.add_heading(header_title, level=1)
    doc.add_paragraph(header_subtitle)
    doc.add_paragraph(f"총 {len(cards)}건의 사례가 포함되어 있습니다.")
    doc.add_paragraph("")
    doc.add_paragraph("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    doc.add_paragraph("【마커 해석 가이드 (전체 카드 공통)】")
    doc.add_paragraph("- 마커 없는 텍스트 = TOBE 검수자 1차 초안 (신뢰도 중)")
    doc.add_paragraph("- <v2>...</v2>     = 현업 합의 부분 (신뢰도 고 ★ 가장 우선해서 모방)")
    doc.add_paragraph("- <v3>...</v3>     = 재개선 진행 중 (신뢰도 잠정 — 직접 인용 금지, 방향성만 참조)")
    doc.add_paragraph("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    doc.add_paragraph("")

    # 각 카드를 paragraph로 추가
    for i, card in enumerate(cards, start=1):
        card["card_index"] = i  # 카드 번호 (해당 docx 내)
        card_text = make_card_text(card)
        for line in card_text.split("\n"):
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"  ✅ {output_path} ({len(cards)}건)")


def extract_goldens1_cards():
    """정답데이터1 메시지검토 시트 31건을 v5.21 양식으로 변환."""
    cards = []
    wb = openpyxl.load_workbook(XLSX_GOLDENS1, data_only=True)
    ws = wb["메시지검토"]

    for r in range(3, ws.max_row + 1):
        구분 = ws.cell(r, 1).value
        팀명 = ws.cell(r, 2).value
        템플릿 = ws.cell(r, 3).value
        초안 = ws.cell(r, 4).value or ''
        검토 = ws.cell(r, 5).value or ''
        최종 = ws.cell(r, 6).value or ''

        # 검토안 또는 최종안이 있어야 카드 가능
        if not (검토 or 최종):
            continue
        if not 템플릿:
            continue

        asis = str(초안).strip() if 초안 else ""
        tobe = str(최종).strip() if 최종 else str(검토).strip()

        if not tobe or len(tobe) < 30:
            continue

        # 팀명 정리 (탭·줄바꿈 제거)
        team_str = str(팀명 or "").replace("\t", "").replace("\n", " ").strip()
        team_clean = team_str.split()[0] if team_str else ""
        # 팀명에서 코드 떼어내기 (예: "상품서비스팀2026-113")
        team_simple = re.sub(r"\d+.*$", "", team_clean).rstrip("-").strip()
        if not team_simple:
            team_simple = team_clean

        domain = classify_team_to_domain(team_simple)
        if domain == "pension":
            sub = classify_pension_sub(템플릿)
            domain_full = f"pension_{sub}" if sub != "misc" else "pension"
        else:
            domain_full = domain

        cards.append({
            "sheet": "메시지검토", "row": r,
            "msg_code": str(구분 or "").strip() if 구분 else "",
            "msg_name": str(템플릿).strip(),
            "team": team_simple,
            "domain": domain_full,
            "label": "HIGH",  # 정답데이터1 = 1차 개선 578건의 대표 발췌, 신뢰도 HIGH 일괄
            "asis_marked": asis,  # 색 정보 없음, 마커 없는 단순 텍스트
            "tobe_marked": tobe,
            "agree": "1차 개선 통과 (회사 공식)",
            "review": "검토완료",
            "feedback_first": "1차 정답지",
            "source": "정답데이터1",
        })

    return cards


def main():
    print("=== 4단계: 통합 카드 변환 시작 ===\n")

    # 1. 정답데이터2 추출
    print("【정답데이터2 추출】")
    cards_v2 = extract_cards_from_xlsx(XLSX_PATH, source_label="정답데이터2")
    print(f"  추출: {len(cards_v2)}건")

    # 2. 중복 제거
    cards_v2_deduped = deduplicate_by_msg_code(cards_v2)
    print(f"  중복 제거 후: {len(cards_v2_deduped)}건 (제거 {len(cards_v2) - len(cards_v2_deduped)}건)")

    # 3. 라벨별 분리
    by_label = defaultdict(list)
    for c in cards_v2_deduped:
        by_label[c["label"]].append(c)
    print(f"  라벨 분포: HIGH={len(by_label['HIGH'])}, MEDIUM={len(by_label['MEDIUM'])}, "
          f"LOW={len(by_label['LOW'])}, NEGATIVE={len(by_label['NEGATIVE'])}, EVAL_ONLY={len(by_label['EVAL_ONLY'])}")

    # vector DB에 들어갈 카드 = HIGH + MEDIUM + LOW
    cards_for_db = by_label["HIGH"] + by_label["MEDIUM"] + by_label["LOW"]
    cards_negative = by_label["NEGATIVE"]
    print(f"  vector DB용: {len(cards_for_db)}건 / NEGATIVE: {len(cards_negative)}건")

    # 4. 도메인별 그룹핑 (positive)
    print("\n【도메인 그룹핑】")
    grouped = domain_grouping(cards_for_db)
    for domain, cs in sorted(grouped.items(), key=lambda x: -len(x[1])):
        print(f"  {domain}: {len(cs)}건")

    # 5. 정답데이터1 추출
    print("\n【정답데이터1 추출】")
    cards_v1 = extract_goldens1_cards()
    print(f"  추출: {len(cards_v1)}건 (모두 HIGH 라벨)")

    # 정답데이터1을 정답데이터2 도메인 그룹에 통합 (v5.21 양식 통일)
    grouped_v1 = domain_grouping(cards_v1)
    print(f"  도메인 분포:")
    for domain, cs in sorted(grouped_v1.items(), key=lambda x: -len(x[1])):
        print(f"    {domain}: {len(cs)}건")

    # 6. docx 작성
    print("\n【카드 docx 작성 — 정답데이터2】")
    output_dir_v2 = Path(r"C:\dev\guide\정답데이터2")
    output_dir_v2.mkdir(exist_ok=True)

    domain_to_filename = {
        "derivatives": "cx_goldens2_derivatives.docx",
        "product": "cx_goldens2_product.docx",
        "credit_loan": "cx_goldens2_credit_loan.docx",
        "settlement": "cx_goldens2_settlement.docx",
        "misc": "cx_goldens2_misc.docx",
        "marketing": "cx_goldens2_marketing.docx",
        "pension": "cx_goldens2_pension.docx",
    }

    for domain, filename in domain_to_filename.items():
        cs = grouped.get(domain, [])
        if not cs:
            print(f"  ⏭️  {filename}: 0건 — skip")
            continue
        output_path = str(output_dir_v2 / filename)
        write_docx(
            cs, output_path,
            header_title=f"미래에셋증권 CX 정답지 (2차 개선 - {domain})",
            header_subtitle=f"회사 공식 2차 메시지 개선 666건 중 {domain} 도메인의 검수 통과 사례 모음. "
                            f"각 사례는 ASIS(원문)와 TOBE(검수자 개선안)를 비교하며, "
                            f"TOBE의 색별 신뢰도(검정/파랑/빨강)는 인라인 마커(<v2>/<v3>)로 표현됩니다.",
        )

    # NEGATIVE
    if cards_negative:
        output_path = str(output_dir_v2 / "cx_exceptions_misuyong.docx")
        write_docx(
            cards_negative, output_path,
            header_title="미래에셋증권 CX 예외 사례 (가이드 룰 적용 자제)",
            header_subtitle="협의결과 = '기존유지' 또는 담당팀 피드백 = '미수용/원안유지' 사례. "
                            "이 메시지 코드/패턴이 등장하면 가이드 룰을 강하게 적용하지 말고 원안에 가깝게 유지할 것. "
                            "사유: 업무 의도 명확화 / 표현 모호성 회피 / 운영 제약.",
        )

    # 정답데이터1 — 31건을 5건씩 6개 docx로 나눠서 덮어쓰기
    print("\n【카드 docx 작성 — 정답데이터1 (덮어쓰기)】")
    output_dir_v1 = Path(r"C:\dev\guide\정답데이터1")
    chunk_size = 5
    suffixes = ["01a", "01b", "01c", "01d", "01e", "01f"]
    for i, suffix in enumerate(suffixes):
        chunk = cards_v1[i * chunk_size:(i + 1) * chunk_size]
        if not chunk:
            break
        filename = f"cx_goldens_{suffix}_review_samples.docx"
        output_path = str(output_dir_v1 / filename)
        write_docx(
            chunk, output_path,
            header_title="미래에셋증권 CX 정답지 (1차 개선 — v5.21 양식)",
            header_subtitle=f"회사 공식 1차 메시지 개선 578건 중 대표 발췌 사례. "
                            f"신뢰도 HIGH (회사 공식 검수 통과). "
                            f"색 정보 없음 (1차 데이터엔 색 라벨링 미진행). "
                            f"이 파일은 {suffix} 묶음 ({len(chunk)}건)입니다.",
        )

    print("\n=== 4단계 완료 ===")


if __name__ == "__main__":
    main()
