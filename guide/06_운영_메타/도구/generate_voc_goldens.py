"""
VOC 골든 카드 4개 신규 작성 — 가이드 PDF p123-126 표준 답변 사례
- p123 기본형 (cx_voc_goldens_basic_p01.docx)
- p124 처리불가 (cx_voc_goldens_reject_p01.docx)
- p125 사실정정 (cx_voc_goldens_misclaim_p01.docx)
- p126 답변편집 (cx_voc_goldens_edit_p01.docx)

VOC 정답데이터 부재 (금/은PC 공통 0.5%) 보완.
LMS 골든 카드 95% 편향 → VOC 입력 시 RAG가 LMS 카드만 반환 → LMS 톤 모방 발생.
이 4개 카드 추가로 VOC 채널 RAG 검색 시 VOC 형식 카드 가용.
"""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt

OUT_DIR = r'C:\dev\guide\output\정답데이터2'

CARDS = [
    {
        'fn': 'cx_voc_goldens_basic_p01.docx',
        'title': 'VOC 표준 답변 정답지 — 기본형 (가이드 123p)',
        'scenario': '콜센터·지점 불친절 등 회사 귀책 X 일반 민원',
        'guide_page': '123p',
        'subtype': 'VOC 기본형',
        'persona': '국내 증권사의 VOC 파트 담당자로서 고객 응대 메시지를 작성',
        'minwon': '''OO지점을 방문할 때마다 고객 응대가 불만족스럽고, 최근에는 상담을 받을 때 윽박지르는 듯한 경험을 했다. 매우 불쾌했다. 기계적으로 응대하지 않았으면 좋겠고 조금 더 친절하게 응대했으면 좋겠다.''',
        'conditions': [
            '도입 문구: 고객 이름+고객님 호명 후 인사 (예: 김미래 고객님)',
            '내용 구성: 고객 감정/상황 공감 → 조치방안 안내 → 재발 방지 의지 표명 → 클로징 순서',
            '작성 분량: 400자 이내 (공백 포함)',
            '전체 톤: 공손·신뢰 (하십시오체 적용)',
            '고객 관심·참여에 감사 인사 + 적극 수용태도',
            '분량: 3~7문장 이내',
        ],
        'forbidden': [
            '"고객님 요청이 처리되었습니다. 확인 후 필요한 조치를 진행하겠습니다." 같은 무심·기계적 말투',
        ],
        'standard_answer': '''김미래 고객님, 안녕하십니까. 지점 방문 과정에서 응대가 만족스럽지 못했고, 상담 중 위압적으로 느껴지셨다는 점에 대해 깊이 공감하며 불편을 드린 점 진심으로 사과드립니다.

고객님께서 말씀 주신 사항은 해당 지점에 즉시 전달하여 응대 태도와 상담 방식 전반을 점검하고, 고객 존중을 기반으로 한 응대 교육을 강화하겠습니다. 앞으로는 기계적인 안내가 아닌, 고객님의 상황을 충분히 경청하고 배려하는 상담이 이루어질 수 있도록 재발 방지에 최선을 다하겠습니다.

소중한 의견을 통해 개선의 기회를 주신 점에 감사드리며, 다시 한 번 신뢰를 회복할 수 있도록 지속적으로 노력하겠습니다.''',
        'key_lessons': [
            '말머리 [미래에셋증권] 사용 안 함',
            '■ 항목 구조 사용 안 함 (자연어 단락 3개)',
            '5단 구조 자발 추가 금지',
            '"꼭 확인해 주세요" 블록 사용 금지',
            '하십시오체 일관 사용',
            '사과 표현 1회 (반복 금지)',
        ],
    },
    {
        'fn': 'cx_voc_goldens_reject_p01.docx',
        'title': 'VOC 표준 답변 정답지 — 처리불가 시나리오 (가이드 124p)',
        'scenario': '해결 불가능한 사항 (시스템 제약/물리적 한계 등) 답변',
        'guide_page': '124p',
        'subtype': 'VOC_reject',
        'persona': '국내 증권사의 VOC 파트 담당자로서 고객 응대 메시지를 작성',
        'minwon': '''대주 신청을 했는데 거래도 안 되고 물량도 부족해서 신청을 못해 속상함''',
        'situation_check': '''대주 예약 주문은 시스템상으로 실행할 수 없다. 물량은 당사에서 미리 알 수 없어서 대응 불가''',
        'conditions': [
            '도입 문구: 고객 이름+고객님 호명 후 인사',
            '내용 구성: 공감 → 사실 확인/원인 결과 안내 → 재발 방지 의지 → 클로징',
            '상황 확인 결과 반영',
            '상황 설명: 이해하기 쉽고 간결 (인과관계 중요)',
            '해결 불가 안내 + 정중한 거절',
            '고객 관심·참여 감사 + 수용태도',
        ],
        'forbidden': [
            '"안 됩니다", "불가능합니다", "무조건 불가합니다" 같은 직설적 표현',
            '고객에게 책임 전가하는 표현',
        ],
        'standard_answer': '''김미래 고객님, 안녕하세요. 대주 신청 과정에서 거래가 이루어지지 않고 물량도 확보되지 않아 많이 답답하셨을 것으로 생각됩니다.

확인된 내용으로는, 현재 대주 예약 주문은 대주시스템 구조상 실행이 보장되지 않아 운영이 어려우며, 대주 가능 물량 또한 증권금융에서 실시간으로 통합 관리되어 당사에서 미리 확인하거나 확보해 두기 어려운 방식으로 운영되고 있습니다. 이러한 구조적 이유로 요청 주신 방식대로 즉각적인 지원을 드리기에는 제한이 있어 너른 양해를 부탁드립니다.

다만 이용 과정에서 겪으신 불편과 의견은 서비스 개선에 중요한 참고가 되므로 관련 부서와 함께 검토하여 대응 방식을 보완할 수 있도록 노력하겠습니다. 추가로 확인이 필요한 사항이 있으시면 언제든지 말씀해 주시기 바랍니다.''',
        'key_lessons': [
            '말머리 [미래에셋증권] 사용 안 함',
            '■ 항목 구조 사용 안 함 (자연어 단락 3개)',
            '"안 됩니다" 직설 거절 금지 → "양해를 부탁드립니다" 정중 거절',
            '시스템 제약을 인과관계로 자연어 설명',
            '책임 전가 표현 금지',
        ],
    },
    {
        'fn': 'cx_voc_goldens_misclaim_p01.docx',
        'title': 'VOC 표준 답변 정답지 — 사실 다른 주장 시나리오 (가이드 125p)',
        'scenario': '고객 주장이 사실과 다른 경우 — 혼선 해소',
        'guide_page': '125p',
        'subtype': 'VOC_misclaim',
        'persona': '국내 증권사의 VOC 파트 담당자로서 고객 응대 메시지를 작성',
        'minwon': '''온라인으로 계좌를 개설하고 지점에 방문했는데 지점 계좌가 없다면서 돌려보냈다. 상품을 못 팔아서 그런 것처럼 느껴져서 기분이 나쁘다.''',
        'situation_check': '''정책상 온라인 계좌와 지점 개설 계좌는 다른 수수료 체계를 가지고 있음. 지점 계좌가 아닌 경우에는 지점에서 상담이 정확하지 않을 수 있어 제한된 사항만 안내가 이루어짐''',
        'conditions': [
            '도입 문구: 고객 이름+고객님 호명 후 인사',
            '내용 구성: 상황 공감 → 사실 정정·혼선 해소 → 마무리 인사',
            '고객 감정 파악·공감·배려 표현',
            '주장이 사실과 다름을 부드럽게 안내',
            '사실 정정: 이해하기 쉽게 짧은 문장 여러 개',
            '"혼선" 단어 직접 사용 금지',
        ],
        'forbidden': [
            '"고객님의 말은 사실과 다르며" 같은 공격적·단정적 표현',
            '책임 전가 표현',
            '과한 사과',
        ],
        'standard_answer': '''김미래 고객님, 안녕하세요. 지점 방문 시 계좌가 없다는 안내를 받으시고 상품 판매 의도가 아닌지 의문이 드셨을 만큼 불편한 감정을 느끼셨을 상황에 충분히 공감합니다.

확인 결과, 온라인으로 개설한 계좌와 지점에서 개설한 계좌는 수수료 체계가 다르게 운영되고 있습니다. 이로 인해 지점 계좌가 아닌 경우에는 지점에서 제공 가능한 상담 범위가 제한되어 일부 안내만 드릴 수 있습니다.

해당 내용이 충분히 설명되지 않아 의도가 다르게 전달될 여지가 있었던 점은 아쉽게 생각합니다. 앞으로는 계좌 유형에 따른 안내 기준을 보다 명확히 전달하여 불편함이 없도록 하겠습니다.

소중한 의견을 전해 주셔서 감사드리며, 언제든 궁금하신 점이 있으면 도움을 드리겠습니다.''',
        'key_lessons': [
            '말머리 [미래에셋증권] 사용 안 함',
            '■ 항목 구조 사용 안 함 (자연어 단락 4개)',
            '"고객님의 말은 사실과 다르며" 같은 공격적 표현 금지',
            '"혼선" 단어 사용 금지 — "다르게 전달될 여지가 있었던 점" 등으로 우회',
            '짧은 문장 여러 개로 사실 정정',
        ],
    },
    {
        'fn': 'cx_voc_goldens_edit_p01.docx',
        'title': 'VOC 표준 답변 정답지 — 답변 초안 편집 (가이드 126p)',
        'scenario': '직원이 작성한 VOC 답변 초안 편집 — 외부 위험 요소 삭제',
        'guide_page': '126p',
        'subtype': 'VOC_edit',
        'persona': '국내 증권사의 VOC 파트 담당자로서 답변 초안을 편집',
        'edit_conditions_5': [
            '민원 요약 포함',
            '민원 접수일 반드시 포함 (예: 11월 10일)',
            '실제 조치만 작성',
            '진심 사과 1회만',
            '고객 기대 + 회사 의지',
        ],
        'forbidden_3_apology': [
            '거듭', '대단히', '진심으로',
        ],
        'forbidden_3_external': [
            '금감원', '언론', '법적 조치',
        ],
        'rules': [
            '분량: 3~5문장 / 최종 답변 하나만',
            '사과 인사: 최초 1회 (강조의 부사어 지양)',
        ],
        'asis_draft': '''안녕하세요 고객님. 거듭 사과드립니다. 대단히 죄송합니다. 진심으로 양해 부탁드립니다.

저희 미래에셋증권은 고객님의 민원을 진지하게 받아들이고 있으며, 이러한 사항이 금감원에 접수되거나 언론에 보도되거나 법적 조치로 이어지지 않도록 최선을 다하겠습니다.

문제 해결을 위해 노력하겠습니다. 다시 한 번 사과드립니다.''',
        'tobe_edited': '''김미래 고객님, 안녕하세요. 11월 10일 접수해 주신 민원 내용 잘 확인했습니다. 이용 과정에서 불편을 드린 점 사과드립니다.

확인 결과 [구체적 조치사항]을 즉시 적용했으며, 동일 사례가 재발하지 않도록 [구체적 개선방안]을 진행하고 있습니다.

고객님께서 기대하시는 안정적이고 신뢰할 수 있는 서비스를 제공할 수 있도록 지속적으로 노력하겠습니다.''',
        'key_lessons': [
            '말머리 [미래에셋증권] 사용 안 함',
            '■ 항목 구조 사용 안 함',
            '과한 사과 부사어 3종 (거듭/대단히/진심으로) 삭제',
            '외부 위험 요소 3종 (금감원/언론/법적 조치) 답변에 포함 금지',
            '민원 접수일 반드시 명시',
            '사과는 1회만',
        ],
    },
]

def render_doc(card):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # 제목
    p = doc.add_paragraph()
    r = p.add_run(card['title'])
    r.bold = True; r.font.size = Pt(14)

    doc.add_paragraph(f"[가이드 {card['guide_page']} 본문 직접 인용 — Mi-Tone v5.21 RAG 자료]")
    doc.add_paragraph(f"[VOC 채널 전용 표준 답변 정답지 — 회사 공식 검수 통과 형태]")
    doc.add_paragraph('')

    # 시나리오
    p = doc.add_paragraph()
    r = p.add_run('■ 시나리오')
    r.bold = True
    doc.add_paragraph(card['scenario'])
    doc.add_paragraph(f"하위 유형 라벨: {card['subtype']}")
    doc.add_paragraph('')

    # 페르소나
    p = doc.add_paragraph()
    r = p.add_run('■ 페르소나 (가이드 본문)')
    r.bold = True
    doc.add_paragraph(card['persona'])
    doc.add_paragraph('')

    # 민원내용 / 상황 (단, edit 카드는 다른 구조)
    if 'minwon' in card:
        p = doc.add_paragraph()
        r = p.add_run('■ 민원내용 (가이드 본문 예시)')
        r.bold = True
        doc.add_paragraph(card['minwon'])
        doc.add_paragraph('')

    if 'situation_check' in card:
        p = doc.add_paragraph()
        r = p.add_run('■ 상황 확인 결과 (가이드 본문)')
        r.bold = True
        doc.add_paragraph(card['situation_check'])
        doc.add_paragraph('')

    # 작성조건 또는 편집조건
    if 'conditions' in card:
        p = doc.add_paragraph()
        r = p.add_run('■ 작성조건 (가이드 본문)')
        r.bold = True
        for c in card['conditions']:
            doc.add_paragraph(f'- {c}')
        doc.add_paragraph('')

    if 'edit_conditions_5' in card:
        p = doc.add_paragraph()
        r = p.add_run('■ 답변 수정조건 5개 (가이드 126p 본문)')
        r.bold = True
        for i, c in enumerate(card['edit_conditions_5'], 1):
            doc.add_paragraph(f'{i}. {c}')
        doc.add_paragraph('')

        p = doc.add_paragraph()
        r = p.add_run('■ 삭제 대상 — 과한 사과 3종 (가이드 126p)')
        r.bold = True
        for x in card['forbidden_3_apology']:
            doc.add_paragraph(f'- {x}')
        doc.add_paragraph('')

        p = doc.add_paragraph()
        r = p.add_run('■ 삭제 대상 — 외부 위험 요소 3종 (가이드 126p)')
        r.bold = True
        for x in card['forbidden_3_external']:
            doc.add_paragraph(f'- {x}')
        doc.add_paragraph('')

        p = doc.add_paragraph()
        r = p.add_run('■ 작성 규칙')
        r.bold = True
        for x in card['rules']:
            doc.add_paragraph(f'- {x}')
        doc.add_paragraph('')

    # 금지표현
    if 'forbidden' in card:
        p = doc.add_paragraph()
        r = p.add_run('■ 금지 표현 (가이드 본문)')
        r.bold = True
        for f in card['forbidden']:
            doc.add_paragraph(f'- {f}')
        doc.add_paragraph('')

    # 표준답변
    if 'standard_answer' in card:
        p = doc.add_paragraph()
        r = p.add_run('━━━ 표준 답변 (가이드 PDF 본문 1:1 인용) ━━━')
        r.bold = True
        doc.add_paragraph('')
        for line in card['standard_answer'].split('\n'):
            doc.add_paragraph(line)
        doc.add_paragraph('')

    # ASIS / TOBE (edit 카드)
    if 'asis_draft' in card:
        p = doc.add_paragraph()
        r = p.add_run('━━━ ASIS (직원 작성 초안 — 편집 전) ━━━')
        r.bold = True
        doc.add_paragraph('')
        for line in card['asis_draft'].split('\n'):
            doc.add_paragraph(line)
        doc.add_paragraph('')

    if 'tobe_edited' in card:
        p = doc.add_paragraph()
        r = p.add_run('━━━ TOBE (편집 후 — 검수자 정정안) ━━━')
        r.bold = True
        doc.add_paragraph('')
        for line in card['tobe_edited'].split('\n'):
            doc.add_paragraph(line)
        doc.add_paragraph('')

    # 핵심 학습 포인트 (Mi-Tone에 가장 중요)
    p = doc.add_paragraph()
    r = p.add_run('━━━ Mi-Tone 학습 포인트 (VOC 형식 모방 시 필수) ━━━')
    r.bold = True
    for k in card['key_lessons']:
        doc.add_paragraph(f'- {k}')
    doc.add_paragraph('')

    # 메타
    p = doc.add_paragraph()
    r = p.add_run('━━━ 메타데이터 ━━━')
    r.bold = True
    doc.add_paragraph(f'- 도메인: VOC')
    doc.add_paragraph(f'- 하위 유형: {card["subtype"]}')
    doc.add_paragraph(f'- 가이드 출처: {card["guide_page"]}')
    doc.add_paragraph(f'- 신뢰도 라벨: HIGH (가이드 PDF 본문 직접 인용 + 회사 공식 표준)')
    doc.add_paragraph(f'- 채널: VOC 답변 (말머리 X, ■구조 X, 자연어 단락)')
    doc.add_paragraph(f'- 형식 핵심: LMS 5단 구조 비적용')

    return doc

results = []
for card in CARDS:
    doc = render_doc(card)
    path = os.path.join(OUT_DIR, card['fn'])
    doc.save(path)
    sz = os.path.getsize(path)
    results.append((card['fn'], sz))

with open(r'C:\dev\guide\_voc_goldens_result.txt', 'w', encoding='utf-8') as f:
    for fn, sz in results:
        f.write(f'{fn}: {sz/1024:.1f}KB\n')
