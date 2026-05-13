"""
v5.21 신규 RAG docx 4종 생성 (PDF 본문 1:1 인용 + 메타 해설 분리)

대상:
1. cx_part4_prompt_meta.docx (가이드 119-128p)
2. cx_editorial_policy.docx (가이드 145-146p)
3. cx_voc_scenario_misclaim.docx (가이드 125p)
4. cx_voc_edit_initial_draft.docx (가이드 126p)

원칙:
- PDF 본문 1:1 인용 (자의 가공 0)
- 헤더 메타에 출처·핵심 룰 요약만 추가
- Mi-Tone 내부 용어 (v5.21 다이어트·tie-breaker 등) 박지 않음
"""

import re
from docx import Document
from pathlib import Path

PDF_DIR = Path(r"C:\dev\guide\_tmp_pdfpages")
OUT_DIR = Path(r"C:\dev\guide\output\guide_rag_cx_docx_flat")

# XML 호환되지 않는 control character 제거 (탭·줄바꿈은 유지)
CTRL_CHAR_RE = re.compile(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]")


def clean(text):
    return CTRL_CHAR_RE.sub("", text)


def read_text(path):
    with open(path, "r", encoding="utf-8") as f:
        return clean(f.read())


def make_docx(title, meta_lines, pdf_text, output_filename):
    doc = Document()
    doc.add_heading(clean(title), 0)
    for line in meta_lines:
        doc.add_paragraph(clean(line))
    doc.add_paragraph("")
    doc.add_paragraph("━" * 40)
    doc.add_heading("가이드 PDF 본문 (1:1 인용)", level=1)
    for line in pdf_text.split("\n"):
        cleaned = clean(line)
        if cleaned.strip():
            doc.add_paragraph(cleaned)
    output_path = OUT_DIR / output_filename
    doc.save(str(output_path))
    print(f"  OK: {output_path}")


def extract_page_block(text, start_marker, end_marker=None):
    """페이지 분리자 사이 본문 추출"""
    if start_marker not in text:
        return ""
    after = text.split(start_marker, 1)[1]
    if end_marker and end_marker in after:
        return after.split(end_marker, 1)[0]
    return after


def main():
    print("=== v5.21 신규 RAG docx 4종 생성 ===\n")

    # 1: cx_part4_prompt_meta (119-128p)
    pdf119_128 = read_text(PDF_DIR / "p119-128-utf8new.txt")
    make_docx(
        "가이드 Part 4 - 챗GPT 프롬프트 작성 가이드 (119-128p)",
        [
            "[가이드 본문 직접 인용 - Mi-Tone v5.21 RAG 자료]",
            "",
            "회사 공식 LLM 활용 메타 가이드.",
            "프롬프트 5 요소: 페르소나 + 정보 + 작성조건 + 제한사항 + 출력형식",
            "",
            "주요 페이지:",
            "- 119p 프롬프트 구성 이해 (이벤트 카피 작성)",
            "- 120p 구체성의 딜레마 ('과도한 장황한 구체성이 LLM 출력을 망친다')",
            "- 121p LMS 인사말 문구 작성 (해요체)",
            "- 122p 편집 프롬프트 (브랜드 보이스 적용)",
            "- 123-126p VOC 4종 답변 (기본형/처리불가/사실다른주장/답변편집)",
            "- 127p 세대별 마케팅 메시지 (Don't 보다 Do 원칙 중심)",
            "",
            "* 가이드 페르소나 예시 (119p, 123p, 127p): \"국내 증권사의 브랜드 마케터/카피라이터/VOC 파트 담당자/마케팅 담당자\" - 이는 ChatGPT 사용자가 ChatGPT에 부여하는 페르소나 예시이며, Mi-Tone 시스템 자체의 페르소나가 아님",
        ],
        pdf119_128,
        "cx_part4_prompt_meta.docx",
    )

    # 2: cx_editorial_policy (145-146p)
    pdf145_146 = read_text(PDF_DIR / "p145-146-utf8.txt")
    make_docx(
        "가이드 Appendix - 미래에셋증권 편집정책 (145-146p)",
        [
            "[가이드 본문 직접 인용 - Mi-Tone v5.21 RAG 자료]",
            "",
            "회사 공식 콘텐츠 작성 원칙. 145-146p Appendix Editorial Policies.",
            "",
            "3대 가치 (145p 본문 그대로):",
            "- 공정성: 특정 자산, 상품, 지역, 기업에 대한 편향 없이 균형 잡힌 시각을 제공합니다.",
            "- 정확성: 모든 수치와 인용은 출처를 명확히 밝히며, 사실에 기반한 정보를 제공합니다.",
            "- 독립성: 외부 광고, 영업, 상품 이해관계로부터 분리되어 독립적인 편집 판단을 유지합니다.",
            "",
            "브랜드 보이스 4종 (145p 본문 그대로):",
            "- Client First | 고객의 상황과 감정을 먼저 헤아립니다.",
            "- Insightful | 전문성을 바탕으로 명료하게 씁니다.",
            "- Trustworthy | 신뢰를 주는 언어로 표현합니다.",
            "- Contributive | 실천 가능한 책임만 진정성 있게 전합니다.",
        ],
        pdf145_146,
        "cx_editorial_policy.docx",
    )

    # 3: cx_voc_scenario_misclaim (125p)
    p125 = extract_page_block(
        pdf119_128, "========== PAGE 125 ==========", "========== PAGE 126 =========="
    )
    make_docx(
        "가이드 Part 4 - VOC 사실 다른 주장 시나리오 (125p)",
        [
            "[가이드 본문 직접 인용 - Mi-Tone v5.21 RAG 자료]",
            "",
            "VOC 4 시나리오 중 '사실 다른 주장' 작성 가이드. 가이드 125p.",
            "",
            "핵심 룰 (125p 본문):",
            "- 작성 포인트: 고객 주장이 사실과 다름 + 혼선 해소",
            "- 내용 구성: 상황 공감 > 사실 정정 및 혼선 해소 > 마무리 인사 순서",
            "- 도입: 고객 이름 + 고객님 호명 후 인사",
            "- 금지 표현: 공격적·단정적·과한 사과·'고객님의 말은 사실과 다르며'",
            "- '혼선' 단어 직접 사용 금지",
            "",
            "표준 답변 사례 (125p): 온라인/지점 계좌 수수료 체계 차이",
        ],
        p125,
        "cx_voc_scenario_misclaim.docx",
    )

    # 4: cx_voc_edit_initial_draft (126p)
    p126 = extract_page_block(
        pdf119_128, "========== PAGE 126 ==========", "========== PAGE 127 =========="
    )
    make_docx(
        "가이드 Part 4 - VOC 답변 초안 편집 시나리오 (126p)",
        [
            "[가이드 본문 직접 인용 - Mi-Tone v5.21 RAG 자료]",
            "",
            "VOC 4 시나리오 중 '답변 초안 편집' 작성 가이드. 가이드 126p.",
            "",
            "답변 수정조건 5개 (126p 본문):",
            "1. 민원 요약 포함",
            "2. 민원 접수일 반드시 포함 (예: 11월 10일)",
            "3. 실제 조치만 작성",
            "4. 진심 사과 (1회만)",
            "5. 고객 기대 + 회사 의지",
            "",
            "삭제 대상 (126p 본문):",
            "- 과한 사과 3종: 거듭 / 대단히 / 진심으로",
            "- 외부 위험 요소 3종: 금감원 / 언론 / 법적 조치",
            "",
            "분량: 3~5문장 / 하나의 최종 답변만",
            "사과 인사는 최초 한 번만 사용 (강조의 부사어 지양)",
        ],
        p126,
        "cx_voc_edit_initial_draft.docx",
    )

    print("\n=== 4 docx 생성 완료 ===")


if __name__ == "__main__":
    main()
